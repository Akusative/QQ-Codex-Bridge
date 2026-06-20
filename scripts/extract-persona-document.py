import io
import sys


def extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            blocks.append(value)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                blocks.append("\t".join(values))
    return "\n".join(blocks)


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def main() -> None:
    extension = sys.argv[1].lower() if len(sys.argv) > 1 else ""
    data = sys.stdin.buffer.read()
    if extension == ".docx":
        text = extract_docx(data)
    elif extension == ".pdf":
        text = extract_pdf(data)
    else:
        raise ValueError("unsupported document type")
    sys.stdout.write(text)


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        sys.stderr.write(f"{type(error).__name__}: {error}\n")
        raise SystemExit(2)
